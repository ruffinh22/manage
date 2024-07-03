#-*- coding: utf-8 -*-
from django.shortcuts import render, get_object_or_404
from operator import itemgetter, attrgetter, methodcaller
from xlutils import copy # http://pypi.python.org/pypi/xlutils
from xlrd import open_workbook # http://pypi.python.org/pypi/xlrd
from xlwt import easyxf
from datetime import datetime
import xlwt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from Etudiant.models import Etu,Appartient
from UE.models import UE
from Note.models import Resultat_UE
from Note.models import Resultat_Semestre
from Matiere.models import Matiere
from Semestre.models import Semestre, InstanceSemestre
from Note.models import Note
from UE.forms import SelectSemestre
from Semestre.forms import SelectInstanceSemestre
# Create your views here.

class etudiant:
	def __init__ (self, numero,resultat):
	#Definition des attributs de chaque instance
		self.numero = numero
		self.resultat = resultat

class UECapitalise:
	def __init__ (self, numero,nom,prenom,ue):
	#Definition des attributs de chaque instance
		self.numero = numero
		self.nom = nom
		self.prenom = prenom
		self.ue = ue

class ResSem:
	def __init__ (self, numero,nom,prenom,resultat, sem1, sem2):
	#Definition des attributs de chaque instance
		self.numero = numero
		self.nom = nom
		self.prenom = prenom
		self.resultat = resultat
		self.sem1 = sem1
		self.sem2 = sem2


def testValidite(string, num):
	if string == "VAL":
		note=0*num
	elif string == "VALC":
		note=1*num
	elif string == "ADAC":
		note=2*num
	elif string == "NATT":
		note=3*num
	elif string == "NATB":
		note=4*num
	elif string == "probleme":
			note=4*num
	else:
		note=5*num
	return note


def ordreListe(semestrePrec, semestre):
	liste = []
	etus = Etu.objects.all()
	for etu in etus:
		note1 = ""
		note2 = ""
		str = 0
		print("Ordre liste")
		print(semestre)
		instanceSemestres = InstanceSemestre.objects.all()
		for ins in instanceSemestres:
			print(ins.id)
		semestre1 = InstanceSemestre.objects.filter(id=semestrePrec)
		semestre2 = InstanceSemestre.objects.filter(id=semestre)
		print("Ordre liste")
		try:
			note1 = Resultat_Semestre.objects.get(etudiant = etu, instance_semestre=semestre1)
			str = testValidite(note1.resultat, 1000)
			note2 = Resultat_Semestre.objects.get(etudiant = etu, instance_semestre=semestre2)
			str += testValidite(note2.resultat, 100)
			moyG = (note1.note+note2.note)/2
			str-=moyG
		except Resultat_Semestre.DoesNotExist:
			print("probleme")
		if str == 0:
			str = 10000
		e = etudiant(etu.apogee, str)
		liste.append(e)
		liste = sorted(liste, key=attrgetter('resultat'))
	for li in liste:
		print (li.numero, li.resultat)
	return liste


"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										PV Semestre 1									 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""




def generationPV_Semestre1():
	document = Document()
	liste = []
	etus = Etu.objects.all()
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	font.size = Pt(10)
	instSemestre = InstanceSemestre.objects.get(semestre__intitule="Semestre 1")
	semestrePrec = instSemestre.semestre.id
	etus = Appartient.objects.filter(instance_semestre=instSemestre)
	maintenant = datetime.now()
	year = maintenant.year
	month = maintenant.month
	if month<7:
		annee = "Année Universitaire " + str(year-1) + "/" + str(year) 
	else:
		annee = "Année Universitaire " + str(year) + "/" + str(year + 1) 
	document.add_paragraph('IUT DU LIMOUSIN')
	
	paragraph = document.add_paragraph('DEPARTEMENT : ..............')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode('Résultats de la délibération du jury du SEMESTRE 1', 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode(annee, 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	
	paragraph = document.add_paragraph()
	run = paragraph.add_run(unicode('Etudiants ayant validé le Semestre 1 :', 'utf-8'))
	font = run.font
	font.underline = True
	capitale = ""
	for etu in etus:
		etudi = Etu.objects.get(apogee=etu.etudiant.apogee)
		try:
			res = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestre)
		except Resultat_Semestre.DoesNotExist:
			print("probleme")
		if(res.note >= 10):
			document.add_paragraph(etu.etudiant.nom + " " + etu.etudiant.prenom, style='ListBullet3')
		else:
			ues = UE.objects.all().filter(semestre=instSemestre.semestre)
			for ue in ues:
				try:
					noteUE = Resultat_UE.objects.get(etudiant=etudi, ue=ue)
				except Resultat_UE.DoesNotExist:
					print("probleme")
				if noteUE.note >= 10:
					capitale += ue.code_ppn + " - "
			if capitale != "":
				etu = UECapitalise(
					numero=etu.etudiant.apogee,
					nom=etu.etudiant.nom,
					prenom=etu.etudiant.prenom,
					ue = capitale
				)
				liste.append(etu)
		capitale = ""

	paragraph = document.add_paragraph()
	run = paragraph.add_run(unicode("Etudiants n'ayant pas valide le semestre1 mais ayant des UE capitalisees : ", 'utf-8'))
	font = run.font
	font.underline = True

	table = document.add_table(rows=1, cols=3,style="TableGrid")
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'nom'
	hdr_cells[1].text = unicode('prénom', 'utf-8')
	hdr_cells[2].text = 'UE CAPITALISEES'
	for etu in liste:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom), 'utf-8')
		row_cells[1].text = unicode(str(etu.prenom), 'utf-8') 
		row_cells[2].text = unicode(str(etu.ue), 'utf-8')
	
	document.add_paragraph()

	paragraph = document.add_paragraph('Limoges, le ' + str(maintenant.day) + "/" + str(maintenant.month) + "/"+ str(maintenant.year))
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph_format.line_spacing

	document.add_page_break()
	document.save('static/documents/PV/pv_semestre.docx')


"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										PV Semestre 2									 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""






def generationPV_Semestre2():
	document = Document()
	listeUe = []
	listeSem = []
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	font.size = Pt(10)
	instSemestrePrec = InstanceSemestre.objects.get(semestre__intitule="Semestre 1")
	instSemestre = InstanceSemestre.objects.get(semestre__intitule="Semestre 2")
	etus = Appartient.objects.filter(instance_semestre=instSemestre)
	maintenant = datetime.now()
	year = maintenant.year
	month = maintenant.month
	if month<7:
		annee = "Année " + str(year-1) + "/" + str(year) 
	else:
		annee = "Année Universitaire" + str(year) + "/" + str(year + 1) 
	document.add_paragraph('IUT DU LIMOUSIN')
	
	paragraph = document.add_paragraph('DEPARTEMENT : ..............')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode('Résultats de la délibération du jury du semestres 1 et 2 de l\'IUT', 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode(annee, 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	
	paragraph = document.add_paragraph()
	font = run.font
	font.underline = True
	capitale = ""
	listeNulle = []
	for etu in etus:
		res = ""
		etudi = Etu.objects.get(apogee=etu.etudiant.apogee)
		try:
			resSem1 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestrePrec)
			resSem2 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestre)
		except Resultat_Semestre.DoesNotExist:
			print("probleme")
		if(resSem1.resultat == "VAL" ):
			res = "S3"	
		elif (resSem2.resultat == "VAL" ):
			res = "S3"
		else:
			listeNulle.append(etu.etudiant)
		etu = ResSem(
			numero = etu.etudiant.apogee,
			nom = etu.etudiant.nom,
			prenom = etu.etudiant.prenom,
			resultat = res,
			sem1 = resSem1.resultat,
			sem2 = resSem2.resultat
		)
		listeSem.append(etu)

	table = document.add_table(rows=1, cols=4,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = unicode('Nom - Prénom', "utf-8")
	hdr_cells[1].text = unicode('Situation à la rentrée prochaine', 'utf-8')
	hdr_cells[2].text = unicode('Décision S1', "utf-8")
	hdr_cells[3].text = unicode('Décision S2', "utf-8")
	for etu in listeSem:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom) + " " + str(etu.prenom), 'utf-8')
		row_cells[1].text = unicode(str(etu.resultat), 'utf-8') 
		row_cells[2].text = unicode(str(etu.sem1), 'utf-8')
		row_cells[3].text = unicode(str(etu.sem2), 'utf-8')
	
	for etu in listeNulle:
		print(etu)
		ues = UE.objects.all().filter(semestre=instSemestre.semestre)
		for ue in ues:
			try:
				noteUE = Resultat_UE.objects.get(etudiant=etu, ue=ue)
			except Resultat_UE.DoesNotExist:
				print("probleme")
			if noteUE.note >= 10:
				capitale += ue.code_ppn + " - "
		if capitale != "":
			etu = UECapitalise(
				numero=etu.apogee,
				nom=etu.nom,
				prenom=etu.prenom,
				ue = capitale
			)
			listeUe.append(etu)
		capitale = ""

	paragraph = document.add_paragraph()
	paragraph = document.add_paragraph()
	run = paragraph.add_run(unicode("Etudiants n'ayant pas valide le semestre2 mais ayant des UE capitalisees : ", 'utf-8'))
	font = run.font
	font.underline = True

	table = document.add_table(rows=1, cols=3,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Nom'
	hdr_cells[1].text = unicode('Prénom', 'utf-8')
	hdr_cells[2].text = 'UE CAPITALISEES'
	for etu in listeUe:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom), 'utf-8')
		row_cells[1].text = unicode(str(etu.prenom), 'utf-8') 
		row_cells[2].text = unicode(str(etu.ue), 'utf-8')

	document.add_paragraph()

	paragraph = document.add_paragraph('Limoges, le ' + str(maintenant.day) + "/" + str(maintenant.month) + "/"+ str(maintenant.year))
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph_format.line_spacing

	document.add_page_break()
	document.save('static/documents/PV/pv_semestre.docx')



"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										PV Semestre 3									 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""






def generationPV_Semestre3():
	document = Document()
	listeUe = []
	listeSem = []
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	font.size = Pt(10)
	instSemestrePrec = InstanceSemestre.objects.get(semestre__intitule="Semestre 2")
	instSemestre = InstanceSemestre.objects.get(semestre__intitule="Semestre 3")
	etus = Appartient.objects.filter(instance_semestre=instSemestre)
	maintenant = datetime.now()
	year = maintenant.year
	month = maintenant.month
	if month<7:
		annee = "Année Universitaire" + str(year-1) + "/" + str(year) 
	else:
		annee = "Année " + str(year) + "/" + str(year + 1) 
	document.add_paragraph('IUT DU LIMOUSIN')
	
	paragraph = document.add_paragraph('DEPARTEMENT : ..............')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode('Résultats de la délibération du jury du semestres 2 et 3 de l\'IUT', 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode(annee, 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	
	paragraph = document.add_paragraph()
	font = run.font
	font.underline = True
	capitale = ""
	listeNulle = []
	for etu in etus:
		res = ""
		etudi = Etu.objects.get(apogee=etu.etudiant.apogee)
		try:
			resSem1 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestrePrec)
			resSem2 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestre)
		except Resultat_Semestre.DoesNotExist:
			print("probleme")
		if(resSem1.resultat == "VAL" and resSem2.resultat == "VAL"):
			res = "S4"	
		elif (resSem2.resultat == "VALC" ):
			res = "S4"
		else:
			listeNulle.append(etu.etudiant)
		etu = ResSem(
			numero = etu.etudiant.apogee,
			nom = etu.etudiant.nom,
			prenom = etu.etudiant.prenom,
			resultat = res,
			sem1 = resSem1.resultat,
			sem2 = resSem2.resultat
		)
		listeSem.append(etu)

	table = document.add_table(rows=1, cols=4,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = unicode('Nom - Prénom', "utf-8")
	hdr_cells[1].text = unicode('Situation à la rentrée prochaine', 'utf-8')
	hdr_cells[2].text = unicode('Décision S2', "utf-8")
	hdr_cells[3].text = unicode('Décision S3', "utf-8")
	for etu in listeSem:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom) + " " + str(etu.prenom), 'utf-8')
		row_cells[1].text = unicode(str(etu.resultat), 'utf-8') 
		row_cells[2].text = unicode(str(etu.sem1), 'utf-8')
		row_cells[3].text = unicode(str(etu.sem2), 'utf-8')
	
	for etu in listeNulle:
		print(etu)
		ues = UE.objects.all().filter(semestre=instSemestre.semestre)
		for ue in ues:
			try:
				noteUE = Resultat_UE.objects.get(etudiant=etu, ue=ue)
			except Resultat_UE.DoesNotExist:
				print("probleme")
			if noteUE.note >= 10:
				capitale += ue.code_ppn + " - "
		if capitale != "":
			etu = UECapitalise(
				numero=etu.apogee,
				nom=etu.nom,
				prenom=etu.prenom,
				ue = capitale
			)
			listeUe.append(etu)
		capitale = ""

	paragraph = document.add_paragraph()
	paragraph = document.add_paragraph()
	run = paragraph.add_run(unicode("Etudiants n'ayant pas valide le semestre 3 mais ayant des UE capitalisees : ", 'utf-8'))
	font = run.font
	font.underline = True

	table = document.add_table(rows=1, cols=3,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Nom'
	hdr_cells[1].text = unicode('Prénom', 'utf-8')
	hdr_cells[2].text = 'UE CAPITALISEES'
	for etu in listeUe:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom), 'utf-8')
		row_cells[1].text = unicode(str(etu.prenom), 'utf-8') 
		row_cells[2].text = unicode(str(etu.ue), 'utf-8')

	document.add_paragraph()

	paragraph = document.add_paragraph('Limoges, le ' + str(maintenant.day) + "/" + str(maintenant.month) + "/"+ str(maintenant.year))
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph_format.line_spacing

	document.add_page_break()
	document.save('static/documents/PV/pv_semestre.docx')



"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										PV Semestre 4									 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""






def generationPV_Semestre4():
	document = Document()
	listeUe = []
	listeSem = []
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	font.size = Pt(10)
	instSemestrePrec = InstanceSemestre.objects.get(semestre__intitule="Semestre 3")
	instSemestre = InstanceSemestre.objects.get(semestre__intitule="Semestre 4")
	etus = Appartient.objects.filter(instance_semestre=instSemestre)
	maintenant = datetime.now()
	year = maintenant.year
	month = maintenant.month
	if month<7:
		annee = "Année Universitaire" + str(year-1) + "/" + str(year) 
	else:
		annee = "Année " + str(year) + "/" + str(year + 1) 
	document.add_paragraph('IUT DU LIMOUSIN')
	
	paragraph = document.add_paragraph('DEPARTEMENT : ..............')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode('Résultats de la délibération du jury du semestres 3 et 4 de l\'IUT', 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	paragraph = document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run(unicode(annee, 'utf-8'))
	font = run.font
	font.bold = True
	font.underline = True

	
	paragraph = document.add_paragraph()
	font = run.font
	font.underline = True
	capitale = ""
	listeNulle = []
	for etu in etus:
		res = ""
		etudi = Etu.objects.get(apogee=etu.etudiant.apogee)
		try:
			resSem1 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestrePrec)
			resSem2 = Resultat_Semestre.objects.get(etudiant = etu.etudiant, instance_semestre=instSemestre)
		except Resultat_Semestre.DoesNotExist:
			print("probleme")
		if(resSem1.resultat == "VAL" and resSem2.resultat == "VAL"):
			res = "S4"	
		elif (resSem2.resultat == "VALC" ):
			res = "S4"
		else:
			listeNulle.append(etu.etudiant)
		etu = ResSem(
			numero = etu.etudiant.apogee,
			nom = etu.etudiant.nom,
			prenom = etu.etudiant.prenom,
			resultat = res,
			sem1 = resSem1.resultat,
			sem2 = resSem2.resultat
		)
		listeSem.append(etu)

	table = document.add_table(rows=1, cols=4,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = unicode('Nom - Prénom', "utf-8")
	hdr_cells[1].text = unicode('Resultat 2ème année', 'utf-8')
	hdr_cells[2].text = unicode('Décision S3', "utf-8")
	hdr_cells[3].text = unicode('Décision S4', "utf-8")
	for etu in listeSem:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom) + " " + str(etu.prenom), 'utf-8')
		row_cells[1].text = unicode(str(etu.resultat), 'utf-8') 
		row_cells[2].text = unicode(str(etu.sem1), 'utf-8')
		row_cells[3].text = unicode(str(etu.sem2), 'utf-8')
	
	for etu in listeNulle:
		print(etu)
		ues = UE.objects.all().filter(semestre=instSemestre.semestre)
		for ue in ues:
			try:
				noteUE = Resultat_UE.objects.get(etudiant=etu, ue=ue)
			except Resultat_UE.DoesNotExist:
				print("probleme")
			if noteUE.note >= 10:
				capitale += ue.code_ppn + " - "
		if capitale != "":
			etu = UECapitalise(
				numero=etu.apogee,
				nom=etu.nom,
				prenom=etu.prenom,
				ue = capitale
			)
			listeUe.append(etu)
		capitale = ""

	paragraph = document.add_paragraph()
	paragraph = document.add_paragraph()
	run = paragraph.add_run(unicode("Etudiants n'ayant pas valide le semestre 4 mais ayant des UE capitalisees : ", 'utf-8'))
	font = run.font
	font.underline = True

	table = document.add_table(rows=1, cols=3,style="TableGrid")
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Nom'
	hdr_cells[1].text = unicode('Prénom', 'utf-8')
	hdr_cells[2].text = 'UE CAPITALISEES'
	for etu in listeUe:
		row_cells = table.add_row().cells
		row_cells[0].text = unicode(str(etu.nom), 'utf-8')
		row_cells[1].text = unicode(str(etu.prenom), 'utf-8') 
		row_cells[2].text = unicode(str(etu.ue), 'utf-8')

	document.add_paragraph()

	paragraph = document.add_paragraph('Limoges, le ' + str(maintenant.day) + "/" + str(maintenant.month) + "/"+ str(maintenant.year))
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	paragraph_format.line_spacing

	document.add_page_break()
	document.save('static/documents/PV/pv_semestre.docx')




"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										Resultat Semestre 1								 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""


def generationSemestre1():
	style = xlwt.easyxf(' alignment: horizontal center, vertical center; borders: left thin, right thin, top thin, bottom thin;')
	book = open_workbook('docJury/S1.xls',formatting_info=True)
	book.sheet_by_index(0)
	newFeuille = copy(book)
	ligne = 7
	cp =0
	instSemestre = InstanceSemestre.objects.get(semestre__intitule="Semestre 1")
	semestrePrec = instSemestre.semestre.id
	etus = Appartient.objects.filter(instance_semestre=instSemestre)
	for etu in etus:
		etudi = Etu.objects.get(apogee=etu.etudiant.apogee)
		sem1=0
		sem2=0
		colonne = 0
		newFeuille.get_sheet(0).write(ligne,colonne,cp, style)
		colonne +=1
		newFeuille.get_sheet(0).write(ligne,colonne,etudi.apogee, style)
		colonne +=1
		newFeuille.get_sheet(0).write(ligne,colonne,etudi.nom, style)
		colonne +=1
		newFeuille.get_sheet(0).write(ligne,colonne,etudi.prenom, style)
		colonne +=2
		ues = UE.objects.all().filter(semestre=instSemestre.semestre)
		for ue in ues:
			try:
				noteUE = Resultat_UE.objects.get(etudiant=etudi, ue=ue)	
				newFeuille.get_sheet(0).write(ligne,colonne,noteUE.note, style)
				colonne +=1
			except Resultat_UE.DoesNotExist:
				print("probleme")
		try:
			resS = Resultat_Semestre.objects.get(etudiant = etudi, instance_semestre=instSemestre)
			colonne =8
			newFeuille.get_sheet(0).write(ligne,colonne,resS.note, style)
			sem1 = resS.note
			colonne +=1
			newFeuille.get_sheet(0).write(ligne,colonne,resS.resultat, style)
		except Resultat_Semestre.DoesNotExist:
			print ("probleme")
		ligne += 1
		
	nom_fichier = instSemestre.semestre.code_ppn
	newFeuille.save('static/documents/resultat_semestre/res_semestre.xls')




"""////////////////////////////////////////////////////////////////////////////////////////
//																						 //
//										Resultat Semestre 1								 //
//																						 //
////////////////////////////////////////////////////////////////////////////////////////"""

def genererDocuments(request):
	if request.method == 'POST':
		u = InstanceSemestre.objects.all()
		form = SelectInstanceSemestre(request.POST, instanceSemestres=u)
		if form.is_valid() :
			instSemestreId = form.cleaned_data['select']
		instSemestre = InstanceSemestre.objects.get(id=instSemestreId)
		if instSemestre.semestre.intitule == "Semestre 1":
			generationPV_Semestre1()
			generationSemestre1()
			nom_fichier_pv = instSemestre.semestre.code_ppn+" - PV"
			nom_fichier_semestre = instSemestre.semestre.code_ppn+" - Resultats"
			res=True
		else:
			if instSemestre.semestre.intitule == "Semestre 2":
				generationPV_Semestre2()
				instSemestrePrec = InstanceSemestre.objects.get(semestre__intitule="Semestre 1")
				book = open_workbook('docJury/S2.xls',formatting_info=True)	
			elif instSemestre.semestre.intitule == "Semestre 3":
				instSemestrePrec =InstanceSemestre.objects.get(semestre__intitule="Semestre 2")
				book = open_workbook('docJury/S3.xls',formatting_info=True)
			elif instSemestre.semestre.intitule == "Semestre 4":	
				instSemestrePrec = InstanceSemestre.objects.get(semestre__intitule="Semestre 3")
				book = open_workbook('docJury/S4.xls',formatting_info=True)
			style = xlwt.easyxf(' alignment: horizontal center, vertical center; borders: left thin, right thin, top thin, bottom thin;')
			book.sheet_by_index(0)
			newFeuille = copy(book)
			etus = Etu.objects.all()
			ligne = 7
			cp =0
			semestrePrec = instSemestrePrec.semestre.id
			print(instSemestreId)
			liste = ordreListe(semestrePrec, instSemestreId)
			for etu in liste:
				etudi = Etu.objects.get(apogee=etu.numero)
				sem1=0
				sem2=0
				colonne = 0
				newFeuille.get_sheet(0).write(ligne,colonne,cp, style)
				colonne +=1
				newFeuille.get_sheet(0).write(ligne,colonne,etudi.apogee, style)
				colonne +=1
				newFeuille.get_sheet(0).write(ligne,colonne,etudi.nom, style)
				colonne +=1
				newFeuille.get_sheet(0).write(ligne,colonne,etudi.prenom, style)
				colonne +=2
				ues = UE.objects.all().filter(semestre=instSemestrePrec.semestre)
				for ue in ues:
					try:
						noteUE = Resultat_UE.objects.get(etudiant=etudi, ue=ue)	
						newFeuille.get_sheet(0).write(ligne,colonne,noteUE.note, style)
						colonne +=1
					except Resultat_UE.DoesNotExist:
						print("probleme")
				try:
					resS = Resultat_Semestre.objects.get(etudiant = etudi, instance_semestre=instSemestrePrec)
					colonne =8
					newFeuille.get_sheet(0).write(ligne,colonne,resS.note, style)
					sem1 = resS.note
					colonne +=1
					newFeuille.get_sheet(0).write(ligne,colonne,resS.resultat, style)
				except Resultat_Semestre.DoesNotExist:
					print("probleme")
				ues = UE.objects.all().filter(semestre=instSemestre.semestre)
				colonne = 13
				for ue in ues:
					try:
						noteUE = Resultat_UE.objects.get(etudiant=etudi, ue=ue)	
						newFeuille.get_sheet(0).write(ligne,colonne,noteUE.note, style)
						colonne +=1
					except Resultat_UE.DoesNotExist:
						print("probleme")
				try:
					resS2 = Resultat_Semestre.objects.get(etudiant = etudi, instance_semestre=instSemestre)
					colonne =16
					newFeuille.get_sheet(0).write(ligne,colonne,resS2.note, style)
					sem2 = resS2.note
					colonne +=1
					newFeuille.get_sheet(0).write(ligne,colonne,resS2.resultat, style)
				except Resultat_Semestre.DoesNotExist:
					print("probleme")
				colonne = 21
				if sem1>0 and sem2>0:
					moyAn = (sem1+sem2)/2
					newFeuille.get_sheet(0).write(ligne,colonne,moyAn, style)
					colonne +=1
					if resS.resultat == "VAL" and resS2.resultat == "VAL":
						jury = "VAL"
					elif resS.resultat == "VAL" and resS2.resultat == "VALC":
						jury = "VAL"
					elif resS.resultat == "VAL" and resS2.resultat == "VALC" or resS2.resultat == "VAL" and resS.resultat == "VALC":
						jury = "VAL"
					elif resS.resultat == "AJPC" and resS2.resultat == "VAL" or resS2.resultat == "AJPC" and resS.resultat == "VAL":
						jury = "NVAL"
					elif resS.resultat == "NATB" or resS2.resultat == "NATB":
						jury = "NVAL"
					else:
						jury = "NVAL"
					newFeuille.get_sheet(0).write(ligne,colonne,jury, style)
				ligne += 1
				cp+=1
			res=True
			nom_fichier_pv = instSemestre.semestre.code_ppn+" - PV"
			nom_fichier_semestre = instSemestre.semestre.code_ppn+" - Resultats"
			newFeuille.save('static/documents/resultat_semestre/res_semestre.xls')
	else :
		res=False
		semestre = InstanceSemestre.objects.all()
		request.session['doc'] = False
		form = SelectInstanceSemestre(instanceSemestres=semestre)
	return render(request, 'contenu_html/genererDocuments.html', locals())




def classementSemestre(request):
	if request.method == 'POST':
		Instsem =InstanceSemestre.objects.all()
		form = SelectInstanceSemestre(request.POST, instanceSemestres=Instsem)
		if form.is_valid() :
			intSem = form.cleaned_data['select']
			res=True
		etuSems = Appartient.objects.filter(instance_semestre=intSem)
		liste = []
		for etuSem in etuSems:
			try:
				etu = Etu.objects.get(apogee=etuSem.etudiant.apogee)
				resSem = Resultat_Semestre.objects.get(etudiant=etu, instance_semestre=intSem)
				e= etudiant(etu.apogee, resSem.note)
				liste.append(e)
			except Resultat_Semestre.DoesNotExist:
				print("probleme")
				
		liste = sorted(liste, key=attrgetter('resultat'), reverse=True)
		book = xlwt.Workbook()
		worksheet = book.add_sheet("Classement")
		i=0
		for li in liste:
			worksheet.write(i, 0,li.numero)
			worksheet.write(i, 1,li.resultat)
			i+=1
		book.save("Classement.xls")
	else :
		res=False
		Instsem =InstanceSemestre.objects.all()
		form = SelectInstanceSemestre(instanceSemestres=Instsem)
	return render(request, 'contenu_html/classementSemestre.html', locals())




